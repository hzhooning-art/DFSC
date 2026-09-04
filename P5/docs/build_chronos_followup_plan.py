from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).parents[2] / "00_计划与信息" / "Project_Chronos_P6-P10后续研究规划.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "20384E"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F4F6F9"
PALE_GOLD = "FFF4D6"
GOLD = "8A6500"
RED = "9B1C1C"
GREEN = "2F6B45"
WHITE = "FFFFFF"
TEXT = "222222"
MUTED = "666666"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
    if len(table.rows) > 1:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=TEXT, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_paragraph(doc, text="", style=None, bold_lead=None, color=TEXT, after=8, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = keep
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text), size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text), size=10.5)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=DEEP_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_run_font(r, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_stage(doc, code, title, years, question, theory, experiments, gates, outputs):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(f"{code}  {title}"), name="Microsoft YaHei", size=13, bold=True, color=BLUE)
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    set_run_font(meta.add_run(f"建议周期：{years}"), size=9.5, bold=True, color=MUTED)
    add_paragraph(doc, question, bold_lead="核心问题：", after=6)
    add_paragraph(doc, theory, bold_lead="理论任务：", after=6)
    add_paragraph(doc, experiments, bold_lead="实验任务：", after=6)
    add_paragraph(doc, gates, bold_lead="阶段验收：", after=6)
    add_paragraph(doc, outputs, bold_lead="预期产物：", after=10)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DEEP_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[name]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(header.add_run("PROJECT CHRONOS & NON-LOCALITY"), size=8.5, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(footer.add_run("P1-P5 后续研究规划  |  "), size=8.5, color=MUTED)
add_page_field(footer)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("PROJECT CHRONOS & NON-LOCALITY"), size=11, bold=True, color=GOLD)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(10)
set_run_font(p.add_run("P1-P5 完成后的后续研究规划"), size=25, bold=True, color=NAVY)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(26)
set_run_font(p.add_run("从可信可微算子到非局域物理 AI 的闭环科学发现"), size=14, color=DEEP_BLUE)

add_callout(
    doc,
    "规划定位",
    "以 P1-P5 已完成的数值、学习、规范和可辨识性基础为起点，重排 P6-P10：先把“无法辨识”转化为主动实验设计，再逐步扩展到部分共享谱、广义记忆算子、Mori-Zwanzig 神经算子和非局域定律发现。",
    fill=PALE_GOLD,
    accent=GOLD,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(70)
set_run_font(p.add_run("规划周期：2026-2034（滚动评估）"), size=10.5, bold=True, color=MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(p.add_run("版本：2026-08-30"), size=9.5, color=MUTED)
doc.add_page_break()

# 1
doc.add_heading("一、执行摘要", level=1)
add_paragraph(doc, "P1-P5 已经完成了宏大计划所需的第一层基础设施，但实际成果路线与原始编号发生了有价值的偏移：P4 建立了可执行的可靠性规范，P5 建立了有限观测下的记忆阶数可辨识与拒绝判定。两项工作补上了原规划中最薄弱的“结论何时可信”环节。")
add_paragraph(doc, "因此，后续不应立即跳向聚变、深地或细胞输运等宏大应用，也不应继续堆叠相似基准。最具连续性的下一问题是：当数据不足以支持机制阶数时，如何决定下一次在哪里观测、增加哪个通道或改变何种实验条件，使系统以最低代价跨越可辨识边界。")
add_callout(doc, "总主线", "可靠传播 -> 误差控制 -> 结构化学习 -> 组件资格 -> 机制可辨识 -> 主动实验设计 -> 广义记忆表示 -> 非局域定律发现。")

doc.add_heading("二、P1-P5 的实际成果重映射", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
headers = ["阶段", "实际形成的核心能力", "在宏大计划中的位置", "尚未解决的缺口"]
for i, text in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_fill(cell, NAVY)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size=9.2, bold=True, color=WHITE)
rows = [
    ("P1", "可微 Mittag-Leffler/DFSC 传播基元与软件生态", "计算基础层", "更广泛核族与统一谱表示"),
    ("P2", "前向值、参数梯度及矩阵作用的联合误差控制", "可信数值层", "非正规、连续谱与更强理论界"),
    ("P3", "结构先验、神经修正、路由和 no-harm 证据", "结构化学习层", "统一统计物理分解与跨任务规律"),
    ("P4", "后端无关的可执行可靠性与一致性规范", "资格与治理层", "行业级采用和外部独立复现"),
    ("P5", "共享有限记忆阶数的可辨识性、边界与拒绝", "科学推断层", "主动观测、部分共享谱和广义机制"),
]
for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        if len(table.rows) % 2 == 0:
            set_cell_fill(cells[i], PALE_GRAY)
        p = cells[i].paragraphs[0]
        set_run_font(p.add_run(text), size=8.8)
set_table_widths(table, [800, 2800, 1900, 3860])

add_paragraph(doc, "这一重映射意味着原计划中“受约束可学习记忆核”和“更广义的机制发现”仍然是未来任务，但它们应建立在 P4/P5 已形成的可信决策层之上，而不是回到仅比较拟合误差的路线。", after=4)

doc.add_heading("三、长期科学命题与范围边界", level=1)
add_paragraph(doc, "长期命题：能否建立一种结构化非局域算子方法，使模型同时回答三个问题：如何高效表示长历史、哪些结构被有限数据真实支持、下一次观测怎样最有效地减少机制不确定性？")
add_paragraph(doc, "计划不预设所有复杂系统都必须由分数阶模型描述，也不把更低预测误差等同于发现物理机制。有限指数、Mittag-Leffler、连续谱、振荡模态和非线性记忆核应被视为竞争表示；当观测不满足可辨识条件时，明确拒绝仍是合法结论。")
add_bullet(doc, "主应用锚点：高分子、金属与黏弹材料的松弛和滞后响应。")
add_bullet(doc, "第二应用阶梯：多孔介质或地下输运，用于检验空间异质与长时间尺度。")
add_bullet(doc, "远期挑战应用：生物输运、复杂流体或能量系统；只有在方法完成外部复现后进入。")
add_bullet(doc, "暂不承诺：直接建立全领域统一方程、短期覆盖聚变等离子体，或仅凭单一数据集宣称新物理定律。")

# Roadmap overview
doc.add_heading("四、P6-P10 总体路线", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for i, text in enumerate(["阶段", "时间窗口", "关键跃迁", "核心决策门"]):
    set_cell_fill(table.rows[0].cells[i], NAVY)
    p = table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(text), size=9.2, bold=True, color=WHITE)
roadmap = [
    ("P6", "2026-2027", "从拒绝诊断到主动观测设计", "是否以更低预算稳定跨越可辨识边界"),
    ("P7", "2027-2028", "从精确共享到层次化部分共享谱", "是否区分公共、组特异与个体模态"),
    ("P8", "2028-2030", "从有限正实极点到广义受约束记忆算子", "是否获得可解释增量且保持物理约束"),
    ("P9", "2029-2031", "从组合修正到 Mori-Zwanzig 神经算子", "是否在强基线下实现跨系统可迁移优势"),
    ("P10", "2031-2034", "从模型选择到非局域定律发现与平台化", "是否输出可验证方程并获独立复现"),
]
for row in roadmap:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        if len(table.rows) % 2 == 0:
            set_cell_fill(cells[i], PALE_BLUE)
        set_run_font(cells[i].paragraphs[0].add_run(text), size=8.8)
set_table_widths(table, [800, 1500, 3300, 3760])

add_stage(
    doc,
    "P6",
    "拒绝驱动的主动观测与序贯实验设计",
    "12-18 个月",
    "当 P5 返回 UNRESOLVED 时，如何选择下一观测时间、通道、试样、激励或实验时长，以最小新增成本提升阶数分辨能力？",
    "建立有限窗口下的序贯可辨识性指标，把局部 Fisher 信息、投影灵敏度谱、相邻速率分离度和拒绝概率转化为可优化的设计准则；给出停止、继续采样和支持某阶的序贯规则。",
    "先在可控生成器上分离时间点、通道数、重复数和噪声的作用，再在 PVA/铜合金等公开材料数据上进行回放式主动采样；增加至少一个完全独立公开数据任务，比较随机、均匀、信息增益和鲁棒设计。",
    "相同误判约束下，主动方案应相对均匀采样减少至少 25% 的观测预算，或在相同预算下显著降低拒绝率；零控制不得因主动策略提高错误阶数支持；必须保留无法跨越边界的区域。",
    "一篇方法论文、可复现实验设计模块、序贯决策 API、观测预算-可辨识性相图。",
)

add_stage(
    doc,
    "P7",
    "层次化部分共享记忆谱",
    "12-18 个月",
    "多通道或多材料是否真的共享完全相同的衰减率，还是只共享一部分公共模态并具有组特异偏移？",
    "构建公共谱、组级偏移、个体幅值和稀疏特异模态的层次化表示；研究共享强度、通道相关和观测不均衡下的局部可辨识条件与一致拒绝条件。",
    "设置完全共享、部分共享、无共享和错误分组四类控制；在材料批次、冷却器状态或传感器组数据上开展 leave-one-group-out 迁移；与独立拟合、完全共享和黑盒多任务模型比较。",
    "必须在部分共享生成器上同时优于完全共享和独立拟合，并在无共享负控制上拒绝强行合并；若只能提高预测而不能恢复层次结构，则降级为多任务预测方法，不宣称机制发现。",
    "层次化记忆谱论文、分组诊断工具、共享强度边界图和数据分组指南。",
)

add_stage(
    doc,
    "P8",
    "受约束的广义记忆算子与谱表示",
    "18-24 个月",
    "有限正衰减指数何时不足，如何统一复极点、连续谱、分布阶与弱非线性记忆，同时保持因果性、耗散性或正性？",
    "建立核函数、谱测度、Laplace 符号、有限状态实现和离散算子作用之间的统一表示；形成正测度、共轭对和耗散约束下的可微参数化与逼近误差分析。",
    "以正实极点、阻尼振荡、稠密连续谱和弱非线性核为四类受控任务，并在材料主任务中检验广义模型是否改善长时外推；对每类模型使用独立参考求值和 P4 式资格测试。",
    "至少两类广义表示在独立参考下通过值/梯度/约束验证，并在预先声明的真实任务上获得可复现增量；若增量仅来自参数数量，则保留为实验能力而不进入 stable。",
    "统一 kernel/operator registry、理论论文与应用论文各一条候选路线、版本化软件发布。",
)

add_stage(
    doc,
    "P9",
    "广义 Mori-Zwanzig 结构化神经算子",
    "18-30 个月，与 P8 后半段交叠",
    "能否把高维系统的粗粒化动力学分解为局部漂移、受约束记忆核和涨落噪声，并由神经算子学习未解析部分而不破坏稳定性？",
    "从 Mori-Zwanzig 投影出发构造可训练分解，明确确定性记忆、随机涨落和神经闭合项的角色；研究耗散、涨落-耗散关系、因果性和 no-harm 门控。万能逼近结果仅在可证明的函数类与拓扑下陈述。",
    "从中等规模可验证系统开始，如粗粒化分子链、黏弹材料点或多孔介质代理模型；采用强 FNO、DeepONet、状态空间模型和纯时序网络基线，执行匹配预算、多种子和长时/OOD 测试。",
    "至少在两个相关任务族中稳定改善样本效率、长时稳定或 OOD 泛化中的两项；若只在单任务有效，则定位为领域模型，不宣称统一框架。",
    "Mori-Zwanzig 神经算子论文、基准套件、热力学约束组件和跨任务适用性相图。",
)

add_stage(
    doc,
    "P10",
    "非局域定律发现与开放基础设施",
    "24-36 个月",
    "能否从多条件时空数据中提出可解释的非局域控制律，并使该定律在独立实验和不同离散化下继续成立？",
    "结合受约束符号回归、谱表示和可辨识性判据，区分数值等价、预测等价和机制可区分；建立候选定律的复杂度惩罚、稳定性检验、反事实预测和外部验证协议。",
    "只选择一个成熟主领域开展闭环发现：训练数据提出候选律，保留条件完成盲测，第二团队或第二数据源完成复现；同步将 P1-P9 能力整合为模块化平台。",
    "候选定律必须通过盲测、单位与对称性检查、边界条件迁移和独立复现；若无法得到唯一机制，则输出可辨识等价类和下一实验建议，而非强行给出单一方程。",
    "非局域定律发现引擎、跨项目标准数据卡、外部复现记录、综述或观点论文，以及面向社区的 benchmark registry。",
)

# Application ladder
doc.add_heading("五、应用阶梯：先深后广", level=1)
add_paragraph(doc, "宏大计划需要重大应用，但应用顺序必须服从方法成熟度。建议采用三级阶梯，每一级都要求独立数据、可验证观测协议和失败边界。")
add_number(doc, "材料与流变主线：继续深挖 PVA、铜合金、钢、聚合物和黏弹材料，形成跨材料但问题一致的松弛/滞后证据链。")
add_number(doc, "地下与多孔介质迁移：在 P6/P7 稳定后引入非均匀采样、空间异质和多时间尺度，检验部分共享谱和主动观测。")
add_number(doc, "远期高挑战系统：仅在 P8/P9 完成资格测试后进入生物输运、复杂流体或能源系统；聚变等离子体属于合作型远期目标，不应作为近期论文承诺。")
add_callout(doc, "应用选择原则", "每篇论文只设置一个主科学问题和一个外部迁移域。跨领域数据用于检验同一机制或边界，不用于堆砌案例。", fill=PALE_GOLD, accent=GOLD)

# software
doc.add_heading("六、软件生态与学术阵地", level=1)
add_bullet(doc, "统一抽象：kernel、propagator、memory realization、diagnostic、design policy 和 conformance record 使用同一注册表。")
add_bullet(doc, "发布治理：stable/experimental 分级、语义版本、变更日志、CITATION、Zenodo DOI、数据许可证和结果清单持续维护。")
add_bullet(doc, "互操作：优先提供 PyTorch/JAX 接口，再与 NeuralOperator、DeepXDE、PETSc/FEM 或图算子工作流建立最小正式适配。")
add_bullet(doc, "外部验证：至少邀请一个独立使用者复现实例；在此之前不把下载量或 CI 通过等同于社区采用。")
add_bullet(doc, "学术传播：先形成可复用 benchmark 和教程，再考虑专题综述、研讨会或国际合作；“形成学派”是长期结果而非短期 KPI。")

# Decision tree
doc.add_heading("七、阶段成功与失败执行树", level=1)
add_paragraph(doc, "每个阶段均采用三档输出，防止项目因追求正结果而无限调参。")
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for i, text in enumerate(["结果状态", "判定", "后续动作"]):
    set_cell_fill(table.rows[0].cells[i], NAVY)
    p = table.rows[0].cells[i].paragraphs[0]
    set_run_font(p.add_run(text), size=9.2, bold=True, color=WHITE)
states = [
    ("GO", "主要理论与实验门均通过，外部任务方向一致", "冻结方法、扩展独立验证并进入下一阶段"),
    ("CONDITIONAL GO", "局部有效或预测增益成立，但机制恢复/外部迁移不足", "缩小声明，形成边界论文或领域方法，不升级为统一理论"),
    ("NO-GO / REDIRECT", "违反零控制、结果依赖后验调参，或无法优于强基线", "保存负结果，回退到诊断/实验设计问题；停止同一路线堆算力"),
]
for idx, row in enumerate(states):
    cells = table.add_row().cells
    fill = PALE_BLUE if idx == 0 else (PALE_GOLD if idx == 1 else "FCE8E6")
    for i, text in enumerate(row):
        set_cell_fill(cells[i], fill)
        set_run_font(cells[i].paragraphs[0].add_run(text), size=9)
set_table_widths(table, [1600, 3550, 4210])

doc.add_heading("八、未来 12 个月执行计划", level=1)
quarters = [
    ("第 1-2 个月", "冻结 P6 问题、设计变量和零控制；整理 P5 可辨识边界为设计目标；建立文献矩阵。"),
    ("第 3-4 个月", "实现时间点/通道/重复数三类主动策略和均匀/随机对照；完成小规模仿真筛选。"),
    ("第 5-6 个月", "进行分离实验：独立改变噪声、时域、通道数和采样预算；确定序贯停止准则。"),
    ("第 7-8 个月", "在 PVA 与铜合金记录上开展回放式主动采样；引入一个完全独立公开数据任务。"),
    ("第 9-10 个月", "执行盲种子评估、强基线、消融、误差校准与边界分析；根据 GO/CONDITIONAL GO/NO-GO 冻结声明。"),
    ("第 11-12 个月", "完成 P6 软件模块、论文初稿、可复现包和外部复现邀请；同步启动 P7 的部分共享最小探针。"),
]
for period, task in quarters:
    add_paragraph(doc, f"{period}：{task}", bold_lead=f"{period}：", after=5)

doc.add_heading("九、量化里程碑与成果分层", level=1)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for i, text in enumerate(["维度", "最低完成", "强成果", "宏大计划级成果"]):
    set_cell_fill(table.rows[0].cells[i], NAVY)
    set_run_font(table.rows[0].cells[i].paragraphs[0].add_run(text), size=9, bold=True, color=WHITE)
metrics = [
    ("理论", "给出局部诊断或充分条件", "形成统一界、序贯保证或一致性结果", "跨模型类的非局域表示与可辨识理论"),
    ("实验", "一个主任务和一个外部任务", "多条件盲测、强基线和边界相图", "独立团队复现或前瞻实验验证"),
    ("软件", "可运行代码、测试和 DOI", "稳定 API、外部适配和持续基准", "被其他研究项目作为依赖采用"),
    ("科学结论", "明确适用与失败区域", "解释机制与观测设计共同作用", "发现可迁移、可证伪的非局域定律"),
]
for row in metrics:
    cells = table.add_row().cells
    for i, text in enumerate(row):
        if len(table.rows) % 2 == 0:
            set_cell_fill(cells[i], PALE_GRAY)
        set_run_font(cells[i].paragraphs[0].add_run(text), size=8.5)
set_table_widths(table, [1200, 2400, 2700, 3060])

doc.add_heading("十、主要风险与约束", level=1)
add_bullet(doc, "理论跨度过大：优先证明可检验的局部或设计条件结果，不提前承诺全局万能逼近与唯一机制恢复。")
add_bullet(doc, "数据不足：把“无法辨识”视为结果，并将贡献转向实验设计边界，而不是用更复杂模型掩盖信息不足。")
add_bullet(doc, "项目碎片化：每个新阶段必须明确继承上一阶段的接口、理论对象和证据，不以新数据集数量作为创新。")
add_bullet(doc, "投稿与维护冲突：P1-P5 的返修、代码修复和 DOI 版本维护优先级高于同时启动多个新论文。")
add_bullet(doc, "宏大应用过早：没有领域合作者和可信数据前，不直接宣称解决聚变、细胞或地下工程核心难题。")
add_bullet(doc, "软件主导叙事：软件是检验和复现载体，论文贡献必须落在数学条件、建模机制或科学结论上。")

doc.add_heading("十一、最终形成的研究体系", level=1)
add_paragraph(doc, "若 P6-P10 按阶段门推进，最终成果不是若干互不相干的论文，而是一套面向复杂长记忆系统的闭环方法：模型能够可靠计算非局域传播，知道数值何时可信，知道结构化学习何时有益，知道观测是否足以支持机制，并在证据不足时主动建议下一实验。")
add_callout(doc, "终极但可检验的目标", "建立一个能够“表示记忆、验证计算、约束学习、拒绝过度解释并指导新观测”的非局域物理 AI 平台。其学术价值由可证明条件、独立数据和外部复现共同决定，而不是由愿景规模本身决定。", fill=PALE_BLUE, accent=DEEP_BLUE)

doc.core_properties.title = "P1-P5 完成后的后续研究规划"
doc.core_properties.subject = "Project Chronos & Non-Locality: P6-P10 roadmap"
doc.core_properties.author = "P1-P5 Research Team"
doc.core_properties.keywords = "non-Markovian, non-local operator, SciML, memory identifiability, active experiment design"
doc.save(OUT)
print(OUT)
